import os
from datetime import datetime
import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import google.generativeai as genai
from dotenv import load_dotenv
import torch
import numpy as np
import pydicom
import gc
from tqdm import tqdm
from joblib import Parallel, delayed
import multiprocessing as mp
import time
from PIL import Image
import torch.nn.functional as F
import cv2
import kagglehub
from roi_extract import RoiExtractor

# Download the dataset
try:
    path = kagglehub.dataset_download("dangnh0611/rsna-breast-cancer-detection-best-ckpts")
    st.success("Dataset downloaded successfully!")
except Exception as e:
    st.error(f"Error downloading dataset: {str(e)}")
    path = None

# Constants for ROI YOLOX
ROI_YOLOX_ENGINE_PATH = path  # Updated to use downloaded model path
ROI_YOLOX_INPUT_SIZE = [416, 416]  # Updated from your constants
ROI_YOLOX_CONF_THRES = 0.5
ROI_YOLOX_NMS_THRES = 0.9
ROI_AREA_PCT_THRES = 0.04
ROI_YOLOX_HW = [(52, 52), (26, 26), (13, 13)]  # Updated from your constants
ROI_YOLOX_STRIDES = [8, 16, 32]
MODEL_INPUT_SIZE = [2048, 1024]  # Model constants

# Constants for DICOM processing
SUID2HEADER = {
    # Add your SUID to header mappings
}

VOILUT_FUNCS_MAP = {
    # Add your VOILUT functions mapping
}

VOILUT_FUNCS_INV_MAP = {v: k for k, v in VOILUT_FUNCS_MAP.items()}

def min_max_scale(img):
    maxv = img.max()
    minv = img.min()
    if maxv > minv:
        return (img - minv) / (maxv - minv)
    else:
        return img - minv  # ==0

class DicomsdlMetadata:
    def __init__(self, ds):
        self.window_widths = ds.WindowWidth
        self.window_centers = ds.WindowCenter
        if self.window_widths is None or self.window_centers is None:
            self.window_widths = []
            self.window_centers = []
        else:
            try:
                if not isinstance(self.window_widths, list):
                    self.window_widths = [self.window_widths]
                self.window_widths = [float(e) for e in self.window_widths]
                if not isinstance(self.window_centers, list):
                    self.window_centers = [self.window_centers]
                self.window_centers = [float(e) for e in self.window_centers]
            except:
                self.window_widths = []
                self.window_centers = []

        # if nan --> LINEAR
        self.voilut_func = ds.VOILUTFunction
        if self.voilut_func is None:
            self.voilut_func = 'LINEAR'
        else:
            self.voilut_func = str(self.voilut_func).upper()
        self.invert = (ds.PhotometricInterpretation == 'MONOCHROME1')
        assert len(self.window_widths) == len(self.window_centers)

def process_dicom(dcm_path):
    # Read DICOM file using pydicom instead of dicomsdl
    dcm = pydicom.dcmread(dcm_path)
    
    # Convert DICOM to array
    image_array = dcm.pixel_array
    
    # Normalize the image
    if image_array.max() != 0:
        image_array = image_array / image_array.max()
    image_array = (image_array * 255).astype(np.uint8)
    
    # Convert to PIL Image
    image = Image.fromarray(image_array)
    return image

def generate_mammogram_details(mammogram_input, model):
    """Generate mammogram details from either a file upload or PIL Image"""
    try:
        # If input is already a PIL Image, use it directly
        if isinstance(mammogram_input, Image.Image):
            image = mammogram_input
        # If input is a file upload
        else:
            image = Image.open(mammogram_input)
        
        current_date = datetime.now().strftime('%Y-%m-%d')
        
        prompt = f"""Analyze this breast mammogram image and provide a detailed report. Fill in ALL fields based on the information you can extract from the image. If you absolutely cannot determine a piece of information, state 'Unable to determine from the provided mammogram image.' Do not use placeholders. Make educated guesses where possible, but clearly indicate when you're making an assumption. Follow this structure:

**BREAST MAMMOGRAM ANALYSIS REPORT**

**1. MAMMOGRAM TECHNICAL DETAILS:**
- MAMMOGRAM MACHINE USED:
- VIEW (CC/MLO):
- BREAST SIDE (Left/Right):

**2. BREAST COMPOSITION:**
- Breast Density Category (A/B/C/D):
- Tissue Distribution:

**3. MAMMOGRAPHIC FINDINGS:**
- Mass (if present):
  * Size:
  * Shape:
  * Margins:
  * Density:
  * Location:
- Calcifications:
  * Type:
  * Distribution:
  * Location:
- Architectural Distortion:
- Asymmetries:
- Skin Changes:
- Nipple Changes:

**4. ASSOCIATED FEATURES:**
- Skin Retraction:
- Nipple Retraction:
- Skin Thickening:
- Axillary Adenopathy:

**5. ASSESSMENT:**
- BIRADS Category (0-6):
- Impression:
- Recommendations:

Report Date: {current_date}
"""
        
        chat_session = model.start_chat(history=[])
        response = chat_session.send_message([prompt, image])
        return response.text
        
    except Exception as e:
        st.error(f"Error generating mammogram details: {str(e)}")
        return None

def create_doc(report_text, mammogram_image):
    doc = Document()
    doc.add_heading('BREAST MAMMOGRAM ANALYSIS REPORT', 0)
    
    for line in report_text.split("\n"):
        if line.strip() == '':
            continue
        if line.startswith('**') and line.endswith('**'):
            doc.add_heading(line.strip('**'), level=1)
        elif line.startswith('-'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())

    doc.add_heading('Mammogram Image:', level=1)
    image_stream = BytesIO(mammogram_image.getvalue())
    doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def generate_mammogram_details_kaggle(mammogram_image):
    """Process mammogram using Kaggle's best model approach"""
    try:
        # Save uploaded file temporarily
        temp_path = "temp_dicom.dcm"
        with open(temp_path, "wb") as f:
            f.write(mammogram_image.getbuffer())
        
        # Initialize ROI extractor
        roi_detector = RoiExtractor(
            engine_path=ROI_YOLOX_ENGINE_PATH,
            input_size=ROI_YOLOX_INPUT_SIZE,
            num_classes=1,
            conf_thres=ROI_YOLOX_CONF_THRES,
            nms_thres=ROI_YOLOX_NMS_THRES,
            class_agnostic=False,
            area_pct_thres=ROI_AREA_PCT_THRES,
            hw=ROI_YOLOX_HW,
            strides=ROI_YOLOX_STRIDES,
            exp=None
        )
        
        # Process the image
        processed_img, roi_coords = _single_decode_crop_save_sdl(
            roi_detector,
            temp_path,
            'temp_processed.png',
            save_backend='cv2',
            return_roi=True
        )
        
        # Display both original and processed images
        col1, col2 = st.columns(2)
        with col1:
            st.image(processed_img, caption='Processed ROI', use_container_width=True)
        with col2:
            if roi_coords is not None:
                # Draw rectangle on original image
                original_img = cv2.imread('temp_processed.png')
                x0, y0, x1, y1 = roi_coords
                cv2.rectangle(original_img, (x0, y0), (x1, y1), (0, 255, 0), 2)
                st.image(original_img, caption='Original with ROI Box', use_container_width=True)
            else:
                st.write("No ROI detected")
        
        # Clean up temporary file
        os.remove(temp_path)
        
        return "Kaggle model analysis results"  # Replace with actual model inference
        
    except Exception as e:
        return f"Error processing image with Kaggle model: {str(e)}"

# from pydicom's source
def _apply_windowing_np_v1(arr,
                           window_width=None,
                           window_center=None,
                           voi_func='LINEAR',
                           y_min=0,
                           y_max=255):
    assert window_width > 0
    y_range = y_max - y_min
    # float64 needed (default) or just float32 ?
    # arr = arr.astype(np.float64)
    arr = arr.astype(np.float32)

    if voi_func in ['LINEAR', 'LINEAR_EXACT']:
        # PS3.3 C.11.2.1.2.1 and C.11.2.1.3.2
        if voi_func == 'LINEAR':
            if window_width < 1:
                raise ValueError(
                    "The (0028,1051) Window Width must be greater than or "
                    "equal to 1 for a 'LINEAR' windowing operation")
            window_center -= 0.5
            window_width -= 1
        below = arr <= (window_center - window_width / 2)
        above = arr > (window_center + window_width / 2)
        between = np.logical_and(~below, ~above)

        arr[below] = y_min
        arr[above] = y_max
        if between.any():
            arr[between] = ((
                (arr[between] - window_center) / window_width + 0.5) * y_range
                            + y_min)
    elif voi_func == 'SIGMOID':
        arr = y_range / (1 +
                         np.exp(-4 *
                                (arr - window_center) / window_width)) + y_min
    else:
        raise ValueError(
            f"Unsupported (0028,1056) VOI LUT Function value '{voi_func}'")
    return arr

def resize_and_pad(img, input_size=MODEL_INPUT_SIZE):
    input_h, input_w = input_size
    ori_h, ori_w = img.shape[:2]
    ratio = min(input_h / ori_h, input_w / ori_w)
    # resize
    img = F.interpolate(img.view(1, 1, ori_h, ori_w),
                        mode="bilinear",
                        scale_factor=ratio,
                        recompute_scale_factor=True)[0, 0]
    # padding
    padded_img = torch.zeros((input_h, input_w),
                             dtype=img.dtype,
                             device='cuda')
    cur_h, cur_w = img.shape
    y_start = (input_h - cur_h) // 2
    x_start = (input_w - cur_w) // 2
    padded_img[y_start:y_start + cur_h, x_start:x_start + cur_w] = img
    padded_img = padded_img.unsqueeze(-1).expand(-1, -1, 3)
    return padded_img

def save_img_to_file(save_path, img, backend='cv2'):
    file_ext = os.path.basename(save_path).split('.')[-1]
    if backend == 'cv2':
        if img.dtype == np.uint16:
            # https://docs.opencv.org/3.4/d4/da8/group__imgcodecs.html#gabbc7ef1aa2edfaa87772f1202d67e0ce
            assert file_ext in ['png', 'jp2', 'tiff', 'tif']
            cv2.imwrite(save_path, img)
        elif img.dtype == np.uint8:
            cv2.imwrite(save_path, img)
        else:
            raise ValueError(
                '`cv2` backend only support uint8 or uint16 images.')
    elif backend == 'np':
        assert file_ext == 'npy'
        np.save(save_path, img)
    else:
        raise ValueError(f'Unsupported backend `{backend}`.')

def _single_decode_crop_save_sdl(roi_extractor,
                                dcm_path,
                                save_path,
                                save_backend='cv2',
                                index=0,
                                return_roi=False):
    dcm = pydicom.dcmread(dcm_path)
    meta = DicomsdlMetadata(dcm)
    info = dcm.getPixelDataInfo()
    if info['SamplesPerPixel'] != 1:
        raise RuntimeError('SamplesPerPixel != 1')
    else:
        shape = [info['Rows'], info['Cols']]

    ori_dtype = info['dtype']
    img = np.empty(shape, dtype=ori_dtype)
    dcm.copyFrameData(index, img)
    img_torch = torch.from_numpy(img.astype(np.int16)).cuda()

    # YOLOX for ROI extraction
    img_yolox = min_max_scale(img_torch)
    img_yolox = (img_yolox * 255)  # float32
    if meta.invert:
        img_yolox = 255 - img_yolox

    # YOLOX infer
    try:
        xyxy, _area_pct, _conf = roi_extractor.detect_single(img_yolox)
        if xyxy is not None:
            x0, y0, x1, y1 = xyxy
            crop = img_torch[y0:y1, x0:x1]
        else:
            crop = img_torch
            xyxy = None
    except:
        print('ROI extract exception!')
        crop = img_torch
        xyxy = None

    # apply voi lut
    if meta.window_widths:
        crop = apply_windowing(crop,
                            window_width=meta.window_widths[0],
                            window_center=meta.window_centers[0],
                            voi_func=meta.voilut_func,
                            y_min=0,
                            y_max=255,
                            backend='torch')
    else:
        print('No windowing param!')
        crop = min_max_scale(crop)
        crop = crop * 255

    if meta.invert:
        crop = 255 - crop
    crop = resize_and_pad(crop, MODEL_INPUT_SIZE)
    crop = crop.to(torch.uint8)
    crop = crop.cpu().numpy()
    save_img_to_file(save_path, crop, backend=save_backend)

    if return_roi:
        return crop, xyxy
    else:
        return crop

def apply_windowing(arr, window_width, window_center, voi_func='LINEAR', y_min=0, y_max=255, backend='torch'):
    if backend == 'torch':
        y_range = y_max - y_min
        
        if voi_func == 'LINEAR':
            if window_width < 1:
                raise ValueError("Window Width must be greater than or equal to 1")
            window_center -= 0.5
            window_width -= 1
            
        below = arr <= (window_center - window_width / 2)
        above = arr > (window_center + window_width / 2)
        between = ~below & ~above

        result = torch.zeros_like(arr)
        result[below] = y_min
        result[above] = y_max
        result[between] = ((arr[between] - window_center) / window_width + 0.5) * y_range + y_min
        
        return result
    else:
        return _apply_windowing_np_v1(arr.cpu().numpy(), window_width, window_center, voi_func, y_min, y_max)

def initialize_roi_detector():
    """Initialize the ROI detector without TensorRT"""
    try:
        roi_detector = RoiExtractor(
            input_size=ROI_YOLOX_INPUT_SIZE,
            num_classes=1,
            conf_thres=ROI_YOLOX_CONF_THRES,
            nms_thres=ROI_YOLOX_NMS_THRES,
            class_agnostic=False,
            area_pct_thres=ROI_AREA_PCT_THRES
        )
        return roi_detector
    except Exception as e:
        st.error(f"Error initializing ROI detector: {str(e)}")
        return None

def visualize_roi(image, roi_coords, caption="ROI Detection Result"):
    """Visualize the ROI on the image with a bounding box"""
    # Convert to numpy array if it's a PIL Image
    if isinstance(image, Image.Image):
        image = np.array(image)
    
    # Make a copy to draw on
    vis_image = image.copy()
    
    # Ensure image is in the right format for drawing
    if vis_image.dtype != np.uint8:
        vis_image = ((vis_image - vis_image.min()) / (vis_image.max() - vis_image.min()) * 255).astype(np.uint8)
    
    # Convert to RGB if grayscale
    if len(vis_image.shape) == 2:
        vis_image = cv2.cvtColor(vis_image, cv2.COLOR_GRAY2RGB)
    
    # Draw ROI rectangle
    if roi_coords is not None:
        x0, y0, x1, y1 = roi_coords
        cv2.rectangle(vis_image, (x0, y0), (x1, y1), (0, 255, 0), 2)  # Green box
        
        # Add text showing coordinates
        text = f"ROI: ({x0}, {y0}), ({x1}, {y1})"
        cv2.putText(vis_image, text, (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 
                    1, (0, 255, 0), 2)
    
    # Convert back to PIL Image for Streamlit
    return Image.fromarray(vis_image)

def process_single_dicom(roi_detector, dcm_path):
    """Process a single DICOM file with ROI detection and visualization"""
    try:
        # Read DICOM
        dcm = pydicom.dcmread(dcm_path)
        img = dcm.pixel_array
        
        # Ensure image is in uint8 format
        if img.dtype != np.uint8:
            img = ((img - img.min()) / (img.max() - img.min()) * 255).astype(np.uint8)
        
        # Detect ROI
        roi_coords, area_pct, _ = roi_detector.detect_single(img)
        
        # Create visualization
        vis_image = visualize_roi(img, roi_coords, "ROI Detection")
        
        # Create columns for original and ROI images
        col1, col2 = st.columns(2)
        
        with col1:
            st.image(img, caption="Original Image", use_container_width=True)
        
        with col2:
            st.image(vis_image, caption=f"ROI Detection (Area: {area_pct:.2%})", use_container_width=True)
        
        # Crop the image if ROI was detected
        if roi_coords is not None:
            x0, y0, x1, y1 = roi_coords
            cropped_img = img[y0:y1, x0:x1]
            st.image(cropped_img, caption="Cropped ROI", use_container_width=True)
            return Image.fromarray(cropped_img)
        else:
            st.warning("No ROI detected, using full image")
            return Image.fromarray(img)
            
    except Exception as e:
        st.error(f"Error processing DICOM: {str(e)}")
        return None

def generate_kaggle_report(image, model_path):
    """Generate report using Kaggle model"""
    try:
        # TODO: Implement Kaggle model inference
        # For now, return a structured report
        current_date = datetime.now().strftime('%Y-%m-%d')
        return f"""**BREAST MAMMOGRAM ANALYSIS REPORT (Kaggle Model)**

**1. TECHNICAL ANALYSIS:**
- Model Used: RSNA Breast Cancer Detection Model
- Analysis Date: {current_date}

**2. DETECTION RESULTS:**
- Probability Score: [Pending Implementation]
- Classification: [Pending Implementation]

**3. RECOMMENDATIONS:**
Based on the current implementation status, please refer to the Gemini analysis for a complete report.
"""
    except Exception as e:
        st.error(f"Error with Kaggle model analysis: {str(e)}")
        return None

def main():
    # Load environment variables and configure Gemini
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if api_key is None:
        st.error("GEMINI_API_KEY is not set in environment variables")
        return
    
    genai.configure(api_key=api_key)
    
    # Configure Gemini model
    generation_config = {
        "temperature": 0.4,
        "top_p": 1,
        "top_k": 32,
        "max_output_tokens": 4096,
    }
    
    # Initialize Gemini model with the newer version
    gemini_model = genai.GenerativeModel(
        model_name="gemini-1.5-pro",
        generation_config=generation_config,
    )

    # Initialize ROI detector
    roi_detector = initialize_roi_detector()
    if roi_detector:
        st.success("ROI detector initialized successfully!")

    st.title("Breast Mammogram Analysis Using AI Generative Model API")

    # Add model selection dropdown
    model_choice = st.selectbox(
        "Choose Analysis Model",
        ["Google Gemini", "Kaggle Best Mammogram Model"]
    )

    st.header("Upload Image")
    mammogram_file = st.file_uploader("Upload Mammogram Image", type=["dcm", "png", "jpg", "jpeg"])

    if mammogram_file is not None:
        # Process image with ROI detection
        with st.spinner("Processing image..."):
            if mammogram_file.name.lower().endswith('.dcm'):
                processed_image = process_single_dicom(roi_detector, mammogram_file)
            else:
                # Handle regular image files
                image = Image.open(mammogram_file)
                img_array = np.array(image)
                roi_coords, area_pct, _ = roi_detector.detect_single(img_array)
                vis_image = visualize_roi(img_array, roi_coords)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.image(image, caption="Original Image", use_container_width=True)
                with col2:
                    st.image(vis_image, caption=f"ROI Detection (Area: {area_pct:.2%})", use_container_width=True)
                
                if roi_coords is not None:
                    x0, y0, x1, y1 = roi_coords
                    cropped_img = img_array[y0:y1, x0:x1]
                    processed_image = Image.fromarray(cropped_img)
                    st.image(processed_image, caption="Cropped ROI", use_container_width=True)
                else:
                    processed_image = image
                    st.warning("No ROI detected, using full image")

        if st.button("Generate Mammogram Report"):
            with st.spinner("Analyzing Mammogram image..."):
                if model_choice == "Google Gemini":
                    mammogram_details = generate_mammogram_details(processed_image, gemini_model)
                else:  # Kaggle Best Mammogram Model
                    try:
                        path = kagglehub.dataset_download("dangnh0611/rsna-breast-cancer-detection-best-ckpts")
                        kaggle_report = generate_kaggle_report(processed_image, path)
                        
                        if kaggle_report:
                            mammogram_details = kaggle_report
                        else:
                            st.warning("Kaggle model analysis failed - Using Gemini as fallback")
                            mammogram_details = generate_mammogram_details(processed_image, gemini_model)
                    except Exception as e:
                        st.error(f"Error with Kaggle model: {str(e)}")
                        st.info("Falling back to Gemini model")
                        mammogram_details = generate_mammogram_details(processed_image, gemini_model)
                
                if mammogram_details:
                    st.header("Generated Mammogram Report")
                    st.markdown(mammogram_details)
                    st.session_state.mammogram_details = mammogram_details

        # Download report button
        if hasattr(st.session_state, 'mammogram_details'):
            doc_file_stream = create_doc(st.session_state.mammogram_details, mammogram_file)
            st.download_button(
                label="Download Mammogram Report",
                data=doc_file_stream,
                file_name="Mammogram_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == '__main__':
    main()