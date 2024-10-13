import os
from datetime import datetime
import streamlit as st
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Configure the generative model
api_key = os.getenv("GEMINI_API_KEY")
if api_key is None:
    raise ValueError("GEMINI_API_KEY is not set in environment variables")
genai.configure(api_key=api_key)

# Set up the model
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 64,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-pro",
    generation_config=generation_config,
)

def generate_ultrasound_details(ecg_image):
    image = Image.open(ecg_image)
    current_date = datetime.now().strftime('%Y-%m-%d')
    
    prompt = f"""Analyze this breast ultrasound image and provide a detailed report. Fill in ALL fields based on the information you can extract from the image. If you absolutely cannot determine a piece of information, state 'Unable to determine from the provided Ultrasound image.' Do not use placeholders like '[To be filled]'. Make educated guesses where possible, but clearly indicate when you're making an assumption. Follow this structure:

**BREAST ULTRASOUND ANALYSIS REPORT**

**1. ULTRASOUND TECHNICAL DETAILS:**
- Ultrasound Machine Used:
- BI-RADS Category:

**2. ULTRASOUND FINDINGS:**
- View:
- Findings:

**3. INTERPRETATION:**
- Normal or Benign or Malignant:
- Diagnosis/Findings:

**4. CONCLUSION AND RECOMMENDATIONS:**
- Summary:
- Recommendations:
- Caution: This analysis is generated by an AI system (MED360). For accurate diagnosis and treatment, please consult with a qualified healthcare professional. It is critical to emphasize that this interpretation is limited by artifact and the available information. Immediate medical evaluation is crucial.

**5. REPORTING SONOGRAPHER:**
- Name:
- Signature: Unable to provide signature for AI-generated report.
- Date of Report: {current_date}
"""

    chat_session = model.start_chat(history=[])
    response = chat_session.send_message([prompt, image])
    return response.text

def create_doc(report_text, ecg_image):
    doc = Document()
    doc.add_heading('BREAST ULTRASOUND ANALYSIS REPORT', 0)
    
    for line in report_text.split("\n"):
        if line.strip() == '':
            continue
        if line.startswith('**') and line.endswith('**'):
            doc.add_heading(line.strip('**'), level=1)
        elif line.startswith('-'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.strip())

    doc.add_heading('Ultrasound Image:', level=1)
    image_stream = BytesIO(ecg_image.getvalue())
    doc.add_picture(image_stream, width=Inches(6))

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def main():
    st.title("Breast Ultrasound Analysis Using AI Generative Model API")

    st.header("Attachments")
    ultrasound_image = st.file_uploader("Upload Ultrasound Image", type=["png", "jpg", "jpeg"])

    if ultrasound_image is not None:
        st.image(ultrasound_image, caption='Uploaded Breast Ultrasound Image', use_column_width=True)

        if st.button("Generate Breast Ultrasound Report"):
            with st.spinner("Analyzing Breast ultrasound image..."):
                ultrasound_details = generate_ultrasound_details(ultrasound_image)
            st.header("Generated Breast Ultrasound Report")
            st.markdown(ultrasound_details)
            
            # Store the generated report in session state
            st.session_state.ecg_details = ultrasound_details

        # Check if report has been generated
        if hasattr(st.session_state, 'ultrasound_details'):
            doc_file_stream = create_doc(st.session_state.ultrasound_details, ultrasound_image)
            st.download_button(
                label="Download Ultrasound Report",
                data=doc_file_stream,
                file_name="Ultrasound_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == '__main__':
    main()